from docx import Document
from docx.shared import Inches

# Importing tex to speach feature
import pyttsx3

# Defining pyttsx3 as a shorter version function
def speak(text):
    pyttsx3.speak(text)

document = Document()

# Profile picture
document.add_picture(
    "IMG_0892.png",
    width = Inches(2.0)
)

# name, phone number, and email details
speak("Hi! I'm an inteligent Resume building assistant! Let me help you with your Resume! What is your name?")
name = input("What is your name? ")
speak("Hello " + name + " how are you today?")

speak("What is your last name?")
last_name = input("What is your last name? ")

speak("Perfect! Lets move on! What is your phone number? ")
phone_number = input("What is your phone number? ")
speak("And what is your email address? ")
email = input("What is your email address? ")

document.add_paragraph(
    name + " " + last_name + "       |       " + phone_number + "       |       " + email)

# About me
document.add_heading("About me")
speak("Now tell me about yourself? ")
document.add_paragraph(
    input("Tell me about youself? ")
)

# Work experience
document.add_heading("Work experience")
# using p as a short cut for paragraph
p = document.add_paragraph()

speak("Wonderful! Now tell me about your work experience.")
print("Wonderful! Now tell me about your work experience.")
company = input("Enter a company ")
start_date = input("Start Date ")
end_date = input("End Date ")

p.add_run(company + "       ").bold = True
p.add_run(start_date + " - " + end_date + "\n").italic = True
p.add_run("\n")
experience_details = input(
    "Describe your experience at " + company + " "
)
p.add_run(experience_details)

# More work experience
while True:
    has_more_experiences = input(
        "Do you have another work experience? Yes or No "
    )
    if has_more_experiences.lower() == "yes":
        p = document.add_paragraph()

        company = input("Enter a company ")
        start_date = input("Start Date ")
        end_date = input("End Date ")

        p.add_run(company + "       ").bold = True
        p.add_run(start_date + " - " + end_date + "\n").italic = True
        p.add_run("\n")
        experience_details = input(
            "Describe your experience at " + company + " "
        )
        p.add_run(experience_details)
    else:
        break

# Skills
speak("Wonderful! Now tell me about your skills.")
document.add_heading("Skils")
print("Perfect! Now tell us about your skills.")
skill = input("Please input your skill: ")
p = document.add_paragraph(skill)
p.style = 'List Bullet'

while True:
    has_more_skills = input(
        "Do you want to add another skill? Yes or No "
    )
    if has_more_skills.lower() == "yes":
        skill = input("Please enter skill: ")
        p = document.add_paragraph(skill)
        p.style = 'List Bullet'
    else:
        break

speak("Nice job! Your resume is redy to rock!")

# Footer
section = document.sections[0]
footer = section.footer
p = footer.paragraphs[0]
p.text = "This resume is generated using Ivan Kharlashkin's code."

document.save("cv.docx")